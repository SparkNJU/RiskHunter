import requests
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document
import xlsxwriter

items = []  #定义一个空列表，为后期数据的存储做好准备
for i in range(1, 134):  #range区间是“左闭右开”的
    if i == 1:
        url = 'https://www.safe.gov.cn/safe/whxw/index.html'
    else:
        url = f'https://www.safe.gov.cn/safe/whxw/index_{i}.html'  # F-Strings是开头有一个f的字符串文字，使用后包含表达式的大括号将被其值替换，表达式在运行时进行渲染，如果不加上f最后爬取出来的会是空列表
    r = requests.get(url)
    r.encoding = r.apparent_encoding  #防止因编码语言不同出现乱码
    text = r.text
    soup = BeautifulSoup(text, 'html.parser')
    li_list = soup.find('div', class_='list_conr').find_all('li')
    for li in li_list:
        title = li.find('a').text
        href = li.find('a').get('href')[1:]  #获取外汇新闻链接后缀
        child_url = 'https://www.safe.gov.cn/' + href
        date = li.find('dd').text
        item = [title, date, child_url]  #列表内容按照新闻标题、新闻时间、新闻子链接排列
        items.append(item)  #将爬取所有item作为一个整体添加到总列表中

#pandas库保存新闻索引数据
df = pd.DataFrame(items, columns=['标题', '发布时间', '链接'])  #将总列表中的列标题分别命名为标题、发布时间、链接
excel_file = '国家外汇管理局外汇新闻信息.xlsx'
workbook = xlsxwriter.Workbook(excel_file)
worksheet = workbook.add_worksheet()

# 写入列名到 Excel 文件
for col_num, value in enumerate(df.columns):
    worksheet.write(0, col_num, value)

# 写入 DataFrame 中的数据到 Excel 文件
for row_num, row_data in df.iterrows():
    for col_num, value in enumerate(row_data):
        worksheet.write(row_num + 1, col_num, value)  # +1 是因为第一行被用于列名

workbook.close()
print(df)

# docx库保存新闻正文内容
for i in range(2467, len(df)):
    url = df.iloc[i]['链接']
    date = df.iloc[i]['发布时间']
    r = requests.get(url)
    r.encoding = r.apparent_encoding
    text = r.text
    soup = BeautifulSoup(text, 'html.parser')
    title = soup.find('div', class_='detail_tit').text  #获取新闻标题
    title = title.replace('/', '')  #下载前要将文件名中带/号的去掉，因为文件命名规则不能带/号，否则程序会中断
    title = title.replace('|', '')  #下载前要将文件名中带|号的去掉，因为文件命名规则不能带|号，否则程序会中断
    content = soup.find('div', class_='detail_content').text  #获取新闻正文
    doc = Document()
    doc.add_heading(title)  #添加标题
    print("i = " + str(i) + "title = " + title)
    doc.add_paragraph(date)  #添加段落
    doc.add_paragraph(content)  #添加段落
    doc.save(f'{title}.docx')  #保存文档
